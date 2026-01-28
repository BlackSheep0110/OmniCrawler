import re
import os
from datetime import datetime
import docx
from urllib.parse import urlparse, unquote
import config

def show_banner():
    """Displays the program banner with author information."""
    banner = f"""
    {'='*60}
       _____  __  __ _   _ _____  _____ _____       __          __   ______ _____  
      / __  \|  \/  | \ | |_   _|/ ____|  __ \     /\ \        / /  |  ____|  __ \ 
     | |  | | \  / |  \| |  | | | |    | |__) |   /  \ \  /\  / /   | |__  | |__) |
     | |  | | |\/| | . ` |  | | | |    |  _  /   / /\ \ \/  \/ /    |  __| |  _  / 
     | |__| | |  | | |\  | _| |_| |____| | \ \  / ____ \  /\  /     | |____| | \ \ 
      \____/|_|  |_|_| \_||_____|\_____|_|  \_\/_/    \_\/  \/      |______|_|  \_\
    
    🚀 OmniCrawler - Advanced General Purpose Web Scraper
    👤 Created by: Ahmad Salami Far
    🌐 GitHub: https://github.com/ahmadsalamifar
    {'='*60}
    """
    print(banner)

class Logger:
    """Handles logging to console and a text file."""
    def __init__(self):
        # Create log file if not exists
        if not os.path.exists(config.LOG_FILE):
            with open(config.LOG_FILE, 'w', encoding='utf-8') as f:
                f.write(f"--- Log Started at {datetime.now()} ---\n")

    def log(self, message, status="INFO"):
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        entry = f"[{timestamp}] [{status}] {message}"
        
        # 1. Print to Console
        print(entry)
        
        # 2. Append to Log File (Safe Logging)
        try:
            with open(config.LOG_FILE, 'a', encoding='utf-8') as f:
                f.write(entry + "\n")
        except:
            pass # If file write fails, don't crash app

def clean_filename(title):
    """Sanitizes a string to be used as a valid filename."""
    clean = re.sub(r'[\\/*?:"<>|]', "", title)
    return clean.strip()[:100] + ".txt"

def normalize_url(link):
    """Highly robust URL cleaner."""
    if not link:
        return None
        
    try:
        link = unquote(link)
    except:
        pass

    link = link.strip()
    
    # Handle CSV/Text garbage attached to the end
    garbage_chars = [',', ';', '،', '|', '\t', '\n']
    for char in garbage_chars:
        if char in link:
            parts = link.split(char)
            best_part = link
            for p in parts:
                if 'http' in p or 'www' in p or '.' in p:
                    best_part = p
                    break
            link = best_part

    link = link.strip(').,;]"\':[]<> \u200c\r\n')

    if len(link) < 4 or " " in link:
        return None
    
    if not link.startswith(('http://', 'https://')):
        if re.match(r'^[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}(/.*)?$', link):
            return 'https://' + link
        else:
            return None 
            
    return link

def get_docx_content(file_path):
    """Extracts text deeply from a .docx file."""
    content_parts = []
    try:
        doc = docx.Document(file_path)
        for p in doc.paragraphs:
            content_parts.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    content_parts.append(cell.text)
        try:
            rels = doc.part.rels
            for rel in rels.values():
                if "hyperlink" in rel.reltype:
                    content_parts.append(rel.target_ref)
        except:
            pass
        return "\n".join(content_parts)
    except Exception as e:
        print(f"Error reading docx {file_path}: {e}")
        return ""